"""
DOCX Resume parsing functionality.

This module extracts relevant information from DOCX resumes using python-docx,
including name, education, and research interests.
"""

import re
import docx
import logging
import spacy
from parser import ResumeParser  # Import base parser to extend functionality

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load the spaCy NLP model
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("Loaded spaCy NLP model for DOCX parser")
except Exception as e:
    logger.warning(f"Could not load spaCy model: {e}")
    logger.warning("Using spaCy's default English model")
    nlp = spacy.blank("en")

class DocxResumeParser(ResumeParser):
    """
    A class to handle parsing DOCX resumes and extracting structured information.
    Extends the ResumeParser class to maintain consistent interface.
    """
    
    def __init__(self, file_path):
        """
        Initialize the DocxResumeParser with a DOCX file path.
        
        Args:
            file_path (str): Path to the DOCX resume file
        """
        self.file_path = file_path
        self.text = self._extract_text()
        self.sections = self._split_into_sections()
        self.doc = self._load_document()
        self.headings = self._extract_headings()
    
    def _load_document(self):
        """
        Load the DOCX document for structured analysis.
        
        Returns:
            docx.Document: The loaded DOCX document
        """
        try:
            return docx.Document(self.file_path)
        except Exception as e:
            logger.error(f"Error loading DOCX document: {e}")
            return None
    
    def _extract_text(self):
        """
        Extract all text from the DOCX file.
        
        Returns:
            str: The entire text content of the DOCX
        """
        try:
            doc = docx.Document(self.file_path)
            full_text = []
            
            # Extract text from each paragraph
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Also extract text from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text.append(paragraph.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def _extract_headings(self):
        """
        Extract headings and their content from the document.
        
        Returns:
            dict: A dictionary mapping heading text to content
        """
        if not self.doc:
            return {}
        
        headings = {}
        current_heading = None
        current_content = []
        
        # Find paragraphs with heading styles
        for para in self.doc.paragraphs:
            # Check if paragraph has a heading style
            if para.style.name.startswith('Heading'):
                # Save previous heading and content
                if current_heading:
                    headings[current_heading] = '\n'.join(current_content)
                
                # Start new heading
                current_heading = para.text.strip().upper()
                current_content = []
            elif current_heading:
                # Add to current heading's content
                if para.text.strip():
                    current_content.append(para.text)
        
        # Save the last heading and content
        if current_heading:
            headings[current_heading] = '\n'.join(current_content)
        
        return headings
    
    def extract_name(self):
        """
        Extract the candidate's name from the resume.
        
        Returns:
            str: The extracted name or empty string if not found
        """
        # In DOCX files, the name is typically at the beginning and often in a larger font
        if self.doc:
            # Check the first few paragraphs
            for i, para in enumerate(self.doc.paragraphs[:5]):
                if para.text.strip() and i < 3:
                    # Check font size if available (assuming name has larger font)
                    try:
                        if hasattr(para.runs[0], 'font') and hasattr(para.runs[0].font, 'size'):
                            font_size = para.runs[0].font.size
                            if font_size and font_size > 12:  # Larger than standard text
                                return para.text.strip()
                    except (IndexError, AttributeError):
                        pass
                    
                    # Use NLP to check if it's a person's name
                    doc = nlp(para.text)
                    
                    # Look for PERSON entities
                    for ent in doc.ents:
                        if ent.label_ == "PERSON":
                            return ent.text
                    
                    # If no entity found but first paragraph is short, it might be the name
                    if len(para.text.split()) <= 4:
                        return para.text.strip()
        
        # Fall back to base class implementation
        return super().extract_name()
    
    def extract_education(self):
        """
        Extract education information from the resume.
        
        Returns:
            list: A list of dictionaries with education details
        """
        education_info = []
        
        # Try to find education section from extracted headings
        education_content = None
        for heading, content in self.headings.items():
            if "EDUCATION" in heading or "ACADEMIC" in heading or "QUALIFICATIONS" in heading:
                education_content = content
                break
        
        if education_content:
            # Process the structured education content
            paragraphs = education_content.split('\n')
            
            degree_keywords = [
                "PhD", "Ph.D", "Doctor of Philosophy",
                "MS", "M.S.", "Master of Science", "Master's", "Masters", "MA", "M.A.",
                "BS", "B.S.", "Bachelor of Science", "Bachelor's", "Bachelors", "BA", "B.A.",
                "MBA", "M.B.A.", "Master of Business Administration"
            ]
            
            current_education = {}
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                # If we have collected some education info and find a new degree/institution, 
                # save the current one and start a new one
                if current_education and any(keyword in paragraph for keyword in degree_keywords):
                    education_info.append(current_education)
                    current_education = {}
                
                # Try to identify degree
                for keyword in degree_keywords:
                    if keyword in paragraph:
                        current_education['degree'] = keyword
                        
                        # Try to find field of study
                        match = re.search(r"\b" + re.escape(keyword) + r"\b[,\s]+(?:in|of)?\s+([^,\n]+)", paragraph)
                        if match:
                            current_education['field'] = match.group(1).strip()
                        break
                
                # Look for institution name
                university_patterns = [
                    r"(?:^|\n|\s)([A-Z][a-zA-Z\s]+(?:University|College|Institute|School))",
                    r"(?:^|\n|\s)(University of [A-Z][a-zA-Z\s]+)"
                ]
                
                for pattern in university_patterns:
                    match = re.search(pattern, paragraph)
                    if match:
                        current_education['institution'] = match.group(1).strip()
                        break
                
                # Look for year
                year_matches = re.findall(r"\b(19\d{2}|20\d{2})\b", paragraph)
                if year_matches:
                    current_education['year'] = max(int(y) for y in year_matches)
            
            # Don't forget the last one
            if current_education and (current_education.get('degree') or current_education.get('institution')):
                education_info.append(current_education)
        
        # If no education info found, fall back to text-based extraction
        if not education_info:
            education_info = super().extract_education()
        
        return education_info
    
    def extract_research_interests(self):
        """
        Extract research interests from the resume.
        
        Returns:
            list: A list of extracted research interests
        """
        interests = []
        
        # Try to find research interests section from extracted headings
        research_content = None
        for heading, content in self.headings.items():
            if "RESEARCH" in heading or "INTERESTS" in heading:
                research_content = content
                break
        
        if not research_content:
            # Try skills section
            for heading, content in self.headings.items():
                if "SKILLS" in heading:
                    research_content = content
                    break
        
        if research_content:
            # Look for bullet points which are common in DOCX files
            # In DOCX, bullet points might be represented differently than in plain text
            
            # Look for bullet-like patterns
            lines = research_content.split('\n')
            for line in lines:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # Look for lines that start with bullet-like characters
                if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                    interest = line[1:].strip()
                    if interest:
                        interests.append(interest)
                # Look for numbered lists
                elif re.match(r'^\d+\.', line):
                    interest = re.sub(r'^\d+\.', '', line).strip()
                    if interest:
                        interests.append(interest)
                # Sometimes interests are just plain text in a paragraph
                elif len(interests) == 0 and len(line) < 100:  # Arbitrary threshold
                    # Split by commas
                    for part in line.split(','):
                        part = part.strip()
                        if part and part.lower() not in ['and', 'or']:
                            interests.append(part)
        
        # If no interests found, fall back to text-based extraction
        if not interests:
            interests = super().extract_research_interests()
        
        return interests

def main():
    """Simple test for the DocxResumeParser"""
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python docx_parser.py [docx_file_path]")
        return
    
    parser = DocxResumeParser(sys.argv[1])
    parsed_data = parser.parse()
    print(parsed_data)

if __name__ == "__main__":
    main()
